#!/usr/bin/env python3

from __future__ import annotations

import argparse
import html
import re
import sys
from dataclasses import dataclass
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:
    print(
        "Missing dependency: python-docx.\n"
        "Create a virtualenv and install it with:\n"
        "  python3 -m venv /tmp/maxos-docx-venv\n"
        "  source /tmp/maxos-docx-venv/bin/activate\n"
        "  pip install python-docx\n",
        file=sys.stderr,
    )
    raise SystemExit(1) from exc

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape as rl_landscape, portrait as rl_portrait
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import PageBreak, Paragraph as RLParagraph, SimpleDocTemplate, Spacer, Table as RLTable, TableStyle

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


TITLE_COLOR = RGBColor(19, 52, 89)
ACCENT_COLOR = RGBColor(28, 96, 168)
TEXT_COLOR = RGBColor(41, 41, 41)
MUTED_COLOR = RGBColor(102, 102, 102)
TABLE_HEADER_FILL = "D9E7F5"


@dataclass
class Heading:
    level: int
    text: str


@dataclass
class Paragraph:
    text: str


@dataclass
class ListItem:
    ordered: bool
    level: int
    text: str
    number: str | None = None


@dataclass
class Table:
    rows: list[list[str]]


@dataclass
class Formula:
    text: str


TOKEN_BOLD = re.compile(r"(\*\*[^*]+\*\*)")
TOKEN_MIXED = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
UL_RE = re.compile(r"^(\s*)-\s+(.*)$")
OL_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
TABLE_SEPARATOR_RE = re.compile(r"^\|?(?:\s*:?-+:?\s*\|)+\s*$")
MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
WIKI_LINK_RE = re.compile(r"\[\[([^\]]+)\]\]")
LATEX_TEXT_RE = re.compile(r"\\text\{([^}]*)\}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export a markdown note to a styled .docx file.")
    parser.add_argument("input", help="Path to the markdown file.")
    parser.add_argument("output", nargs="?", help="Output .docx path. Defaults to input stem + .docx")
    parser.add_argument("--doc-title", help="Optional document title override.")
    parser.add_argument(
        "--orientation",
        choices=("portrait", "landscape"),
        help="Page orientation for the exported Word document.",
    )
    parser.add_argument("--pdf", action="store_true", help="Also export a companion PDF.")
    parser.add_argument("--pdf-output", help="Optional output PDF path. Defaults to the .docx stem plus .pdf")
    return parser.parse_args()


def read_source(path: Path) -> tuple[dict[str, str], str]:
    text = path.read_text(encoding="utf-8")
    metadata: dict[str, str] = {}
    if text.startswith("---\n"):
        parts = text.split("\n---\n", 1)
        if len(parts) == 2:
            for line in parts[0][4:].splitlines():
                stripped = line.strip()
                if not stripped or stripped.startswith("#") or ":" not in stripped:
                    continue
                key, value = stripped.split(":", 1)
                metadata[key.strip()] = value.strip().strip('"').strip("'")
            return metadata, parts[1].lstrip("\n")
    return metadata, text


def is_truthy(value: str | None) -> bool:
    if value is None:
        return False
    return value.strip().lower() in {"1", "true", "yes", "y", "on"}


def strip_markup(text: str) -> str:
    def replace_md_link(match: re.Match[str]) -> str:
        return match.group(1)

    def replace_wiki_link(match: re.Match[str]) -> str:
        target = match.group(1)
        if "|" in target:
            _, label = target.split("|", 1)
            return label
        label = target.split("/")[-1]
        return label.removesuffix(".md")

    text = MD_LINK_RE.sub(replace_md_link, text)
    text = WIKI_LINK_RE.sub(replace_wiki_link, text)
    return text


def clean_text(text: str) -> str:
    return strip_markup(text).replace("\u00a0", " ").strip()


def clean_formula(text: str) -> str:
    text = LATEX_TEXT_RE.sub(r"\1", text)
    text = text.replace("\\text{-}", "-")
    text = text.replace("\\ ", " ")
    text = text.replace("\\", "")
    return text.strip()


def to_pdf_markup(text: str) -> str:
    text = strip_markup(text)
    parts = TOKEN_MIXED.split(text)
    markup_parts: list[str] = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            markup_parts.append(f"<b>{html.escape(part[2:-2])}</b>")
            continue
        if part.startswith("`") and part.endswith("`"):
            markup_parts.append(f'<font name="Courier">{html.escape(part[1:-1])}</font>')
            continue
        if part.startswith("*") and part.endswith("*"):
            markup_parts.append(f"<i>{html.escape(part[1:-1])}</i>")
            continue
        markup_parts.append(html.escape(part))
    return "".join(markup_parts)


def split_cover_block(lines: list[str]) -> tuple[dict[str, str], list[str]]:
    cover_data: dict[str, str] = {}
    filtered: list[str] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        if line.startswith("## Cover"):
            index += 1
            while index < len(lines):
                current = lines[index]
                if current.startswith("## "):
                    break
                stripped = current.strip()
                if stripped.startswith("- ") and ":" in stripped[2:]:
                    key, value = stripped[2:].split(":", 1)
                    cover_data[key.strip()] = clean_text(value.strip())
                index += 1
            continue
        filtered.append(line)
        index += 1
    return cover_data, filtered


def parse_elements(lines: list[str]) -> list[object]:
    elements: list[object] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped == "$$":
            math_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                math_lines.append(lines[index].rstrip())
                index += 1
            if index < len(lines) and lines[index].strip() == "$$":
                index += 1
            elements.append(Formula(clean_formula(" ".join(math_lines))))
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_text(heading_match.group(2))
            elements.append(Heading(level, text))
            index += 1
            continue

        if "|" in stripped and index + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[index + 1].strip()):
            rows: list[list[str]] = []
            while index < len(lines) and "|" in lines[index]:
                row = [clean_text(cell) for cell in lines[index].strip().strip("|").split("|")]
                rows.append(row)
                index += 1
            rows.pop(1)
            elements.append(Table(rows))
            continue

        ul_match = UL_RE.match(line)
        if ul_match:
            indent = len(ul_match.group(1).replace("\t", "    "))
            level = min(indent // 2, 2)
            elements.append(ListItem(False, level, clean_text(ul_match.group(2))))
            index += 1
            continue

        ol_match = OL_RE.match(line)
        if ol_match:
            indent = len(ol_match.group(1).replace("\t", "    "))
            level = min(indent // 2, 2)
            elements.append(ListItem(True, level, clean_text(ol_match.group(3)), ol_match.group(2)))
            index += 1
            continue

        paragraph_lines = [line.rstrip()]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            candidate_stripped = candidate.strip()
            if not candidate_stripped:
                break
            if candidate_stripped == "$$":
                break
            if HEADING_RE.match(candidate):
                break
            if UL_RE.match(candidate) or OL_RE.match(candidate):
                break
            if "|" in candidate_stripped and index + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[index + 1].strip()):
                break
            paragraph_lines.append(candidate.rstrip())
            index += 1
        elements.append(Paragraph(clean_text(" ".join(paragraph_lines))))
    return elements


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def prepare_document_content(input_path: Path, doc_title_override: str | None) -> tuple[dict[str, str], str, dict[str, str], list[object]]:
    metadata, markdown = read_source(input_path)
    lines = markdown.splitlines()

    title = input_path.stem
    if lines and lines[0].startswith("# "):
        title = clean_text(lines[0][2:])
        lines = lines[1:]
        while lines and not lines[0].strip():
            lines.pop(0)

    doc_title = doc_title_override or title
    cover_data, filtered_lines = split_cover_block(lines)
    elements = parse_elements(filtered_lines)
    return metadata, doc_title, cover_data, elements


def apply_document_defaults(document: Document, doc_title: str, orientation: str) -> None:
    section = document.sections[0]
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.1)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    styles = document.styles
    normal_size = Pt(10 if orientation == "landscape" else 10.5)
    title_size = Pt(24 if orientation == "landscape" else 26)
    heading_1_size = Pt(15 if orientation == "landscape" else 16)

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = normal_size
    normal.font.color.rgb = TEXT_COLOR
    normal.paragraph_format.line_spacing = 1.1 if orientation == "landscape" else 1.15
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = title_size
    title.font.bold = True
    title.font.color.rgb = TITLE_COLOR
    title.paragraph_format.space_after = Pt(10)

    subtitle = styles.add_style("Memo Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.base_style = styles["Normal"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(12.5)
    subtitle.font.color.rgb = MUTED_COLOR
    subtitle.paragraph_format.space_after = Pt(6)

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Aptos Display"
    heading_1.font.size = heading_1_size
    heading_1.font.bold = True
    heading_1.font.color.rgb = TITLE_COLOR
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(8)

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Aptos"
    heading_2.font.size = Pt(12.5)
    heading_2.font.bold = True
    heading_2.font.color.rgb = ACCENT_COLOR
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(4)

    heading_3 = styles["Heading 3"]
    heading_3.font.name = "Aptos"
    heading_3.font.size = Pt(11)
    heading_3.font.bold = True
    heading_3.font.color.rgb = TEXT_COLOR
    heading_3.paragraph_format.space_before = Pt(8)
    heading_3.paragraph_format.space_after = Pt(2)

    formula = styles.add_style("Formula", WD_STYLE_TYPE.PARAGRAPH)
    formula.base_style = styles["Normal"]
    formula.font.name = "Aptos Mono"
    formula.font.size = Pt(10)
    formula.font.color.rgb = TITLE_COLOR
    formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(6)
    formula.paragraph_format.space_after = Pt(8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.style = styles["Normal"]
    footer_run = footer.add_run(f"{doc_title}  |  ")
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = MUTED_COLOR
    add_page_number(footer)


def add_cover_page(document: Document, title: str, cover_data: dict[str, str]) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(title)

    prepared_for = cover_data.get("Prepared for")
    if prepared_for:
        p = document.add_paragraph(style="Memo Subtitle")
        p.add_run("Prepared for: ").bold = True
        p.add_run(prepared_for)

    prepared_by = cover_data.get("Prepared by")
    if prepared_by:
        p = document.add_paragraph(style="Memo Subtitle")
        p.add_run("Prepared by: ").bold = True
        p.add_run(prepared_by)

    date = cover_data.get("Date")
    if date:
        p = document.add_paragraph(style="Memo Subtitle")
        p.add_run("Date: ").bold = True
        p.add_run(date)

    status = cover_data.get("Status")
    if status:
        p = document.add_paragraph(style="Memo Subtitle")
        p.add_run("Status: ").bold = True
        p.add_run(status)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)

    divider = document.add_paragraph()
    divider.paragraph_format.space_after = Pt(18)
    run = divider.add_run(" ")
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1C60A8")
    border.append(bottom)
    divider._p.get_or_add_pPr().append(border)
    run.add_break(WD_BREAK.PAGE)


def export_pdf(doc_title: str, cover_data: dict[str, str], elements: list[object], output_path: Path, orientation: str) -> None:
    if not REPORTLAB_AVAILABLE:
        print(
            "Missing dependency: reportlab.\n"
            "Install it in the export virtualenv with:\n"
            "  source /tmp/maxos-docx-venv/bin/activate\n"
            "  pip install reportlab\n",
            file=sys.stderr,
        )
        raise SystemExit(1)

    page_size = rl_landscape(A4) if orientation == "landscape" else rl_portrait(A4)
    left_right_margin = 1.6 * cm if orientation == "landscape" else 2.2 * cm
    top_margin = 1.6 * cm if orientation == "landscape" else 2.1 * cm
    bottom_margin = 1.5 * cm if orientation == "landscape" else 2.0 * cm
    base_font_size = 10 if orientation == "landscape" else 10.5

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=page_size,
        leftMargin=left_right_margin,
        rightMargin=left_right_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
        title=doc_title,
    )

    sample_styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "MemoNormal",
        parent=sample_styles["BodyText"],
        fontName="Helvetica",
        fontSize=base_font_size,
        leading=13,
        textColor=colors.HexColor("#292929"),
        spaceAfter=6,
    )
    title_style = ParagraphStyle(
        "MemoTitle",
        parent=sample_styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=24 if orientation == "landscape" else 26,
        leading=28,
        textColor=colors.HexColor("#133459"),
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "MemoSubtitle",
        parent=normal,
        fontName="Helvetica",
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#666666"),
        spaceAfter=6,
    )
    heading_1 = ParagraphStyle(
        "MemoHeading1",
        parent=normal,
        fontName="Helvetica-Bold",
        fontSize=15 if orientation == "landscape" else 16,
        leading=19,
        textColor=colors.HexColor("#133459"),
        spaceBefore=18,
        spaceAfter=8,
    )
    heading_2 = ParagraphStyle(
        "MemoHeading2",
        parent=normal,
        fontName="Helvetica-Bold",
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor("#1C60A8"),
        spaceBefore=12,
        spaceAfter=4,
    )
    heading_3 = ParagraphStyle(
        "MemoHeading3",
        parent=normal,
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#292929"),
        spaceBefore=8,
        spaceAfter=2,
    )
    formula_style = ParagraphStyle(
        "MemoFormula",
        parent=normal,
        fontName="Courier",
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor("#133459"),
        alignment=1,
        spaceBefore=6,
        spaceAfter=8,
    )
    table_header_style = ParagraphStyle(
        "MemoTableHeader",
        parent=normal,
        fontName="Helvetica-Bold",
        fontSize=9.2,
        leading=11,
        textColor=colors.HexColor("#133459"),
    )
    table_cell_style = ParagraphStyle(
        "MemoTableCell",
        parent=normal,
        fontName="Helvetica",
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#292929"),
    )

    story: list[object] = []
    story.append(RLParagraph(to_pdf_markup(doc_title), title_style))
    for label in ("Prepared for", "Prepared by", "Date", "Status"):
        value = cover_data.get(label)
        if value:
            story.append(RLParagraph(f"<b>{html.escape(label)}:</b> {to_pdf_markup(value)}", subtitle_style))
    story.append(Spacer(1, 18))
    story.append(PageBreak())

    for element in elements:
        if isinstance(element, Heading):
            if element.level in {1, 2}:
                style = heading_1
            elif element.level == 3:
                style = heading_2
            else:
                style = heading_3
            story.append(RLParagraph(to_pdf_markup(element.text), style))
            continue

        if isinstance(element, Paragraph):
            story.append(RLParagraph(to_pdf_markup(element.text), normal))
            continue

        if isinstance(element, ListItem):
            list_style = ParagraphStyle(
                f"MemoList{int(element.ordered)}{element.level}",
                parent=normal,
                leftIndent=(0.45 + (0.45 * element.level)) * cm,
                firstLineIndent=0,
            )
            prefix = f"{element.number}." if element.ordered and element.number else "•"
            story.append(RLParagraph(f"{html.escape(prefix)} {to_pdf_markup(element.text)}", list_style))
            continue

        if isinstance(element, Formula):
            story.append(RLParagraph(html.escape(element.text), formula_style))
            continue

        if isinstance(element, Table):
            usable_width = document.width
            col_count = len(element.rows[0])
            col_widths = [usable_width / col_count] * col_count
            table_rows = []
            for row_index, row in enumerate(element.rows):
                row_style = table_header_style if row_index == 0 else table_cell_style
                table_rows.append([RLParagraph(to_pdf_markup(cell), row_style) for cell in row])
            table = RLTable(table_rows, colWidths=col_widths, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E7F5")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#133459")),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#C2CBD6")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 6))

    def draw_footer(canvas, pdf_doc) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        footer_text = f"{doc_title}  |  {canvas.getPageNumber()}"
        canvas.drawRightString(pdf_doc.pagesize[0] - pdf_doc.rightMargin, 0.75 * cm, footer_text)
        canvas.restoreState()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def add_inline_runs(paragraph, text: str) -> None:
    text = strip_markup(text)
    parts = TOKEN_MIXED.split(text)
    if not parts:
        paragraph.add_run(text)
        return

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Aptos Mono"
            run.font.size = Pt(9.5)
            continue
        if part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            continue
        paragraph.add_run(part)


def export_markdown(
    doc_title: str,
    cover_data: dict[str, str],
    elements: list[object],
    output_path: Path,
    orientation: str,
) -> None:
    document = Document()
    apply_document_defaults(document, doc_title, orientation)
    add_cover_page(document, doc_title, cover_data)

    for element in elements:
        if isinstance(element, Heading):
            if element.level == 1:
                level = 1
            elif element.level == 2:
                level = 1
            elif element.level == 3:
                level = 2
            else:
                level = 3
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline_runs(paragraph, element.text)
            continue

        if isinstance(element, Paragraph):
            paragraph = document.add_paragraph(style="Normal")
            add_inline_runs(paragraph, element.text)
            continue

        if isinstance(element, ListItem):
            if element.ordered:
                style_name = "List Number" if element.level == 0 else f"List Number {element.level + 1}"
            else:
                style_name = "List Bullet" if element.level == 0 else f"List Bullet {element.level + 1}"
            paragraph = document.add_paragraph(style=style_name)
            add_inline_runs(paragraph, element.text)
            continue

        if isinstance(element, Formula):
            paragraph = document.add_paragraph(style="Formula")
            paragraph.add_run(element.text)
            continue

        if isinstance(element, Table):
            table = document.add_table(rows=len(element.rows), cols=len(element.rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for row_index, row in enumerate(element.rows):
                for col_index, cell_text in enumerate(row):
                    cell = table.cell(row_index, col_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    paragraph.style = document.styles["Normal"]
                    add_inline_runs(paragraph, cell_text)
                    if row_index == 0:
                        set_cell_shading(cell, TABLE_HEADER_FILL)
                        for run in paragraph.runs:
                            run.bold = True
            document.add_paragraph()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve() if args.output else input_path.with_suffix(".docx")
    metadata, doc_title, cover_data, elements = prepare_document_content(input_path, args.doc_title)
    orientation = args.orientation or metadata.get("export_orientation") or "portrait"
    export_markdown(doc_title, cover_data, elements, output_path, orientation)

    generated_paths = [output_path]
    if args.pdf or is_truthy(metadata.get("export_pdf")):
        pdf_path = Path(args.pdf_output).expanduser().resolve() if args.pdf_output else output_path.with_suffix(".pdf")
        export_pdf(doc_title, cover_data, elements, pdf_path, orientation)
        generated_paths.append(pdf_path)

    for path in generated_paths:
        print(path)


if __name__ == "__main__":
    main()